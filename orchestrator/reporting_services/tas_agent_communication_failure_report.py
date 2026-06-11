import os
import sys
import asyncio
import aiohttp
import re
import jinja2                                                          
from datetime import datetime, timedelta, timezone
from collections import defaultdict
import urdhva_base
import hpcl_ceg_model                                                 
import orchestrator.notification_manager.notification_factory as notification_factory  
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IST = timezone(timedelta(hours=5, minutes=30))

THINGSBOARD_URL      = urdhva_base.settings.things_board_url
THINGSBOARD_USERNAME = urdhva_base.settings.things_board_username
THINGSBOARD_PASSWORD = urdhva_base.settings.things_board_password

TELEMETRY_KEY        = "Primary Gauge HIGH"
GAP_THRESHOLD_MIN    = 35
PAGE_SIZE            = 1000
EXCLUDED_LOCATIONS   = {"mathura"}
TANK_DEVICE_EXCEPTIONS: dict[str, str] = {
    "secunderabad": "51-TT-",
}


# ── EMAIL: Fetch recipient list from DB ───────────────────────────────────────
async def get_email_users_by_type(email_type, audience):
    all_users = await hpcl_ceg_model.DailyEmailNotificationUsers.get_all(resp_type='plain')
    to_recipients, cc_recipients, bcc_recipients = [], [], []
    for user in all_users.get("data", []):
        if user.get("audience") == audience and \
                (user.get("email_type") or "").lower() == email_type.lower():
            to_recipients.extend(user.get("to_recipients", []))
            cc_recipients.extend(user.get("cc_recipients", []))
            bcc_recipients.extend(user.get("bcc_recipients", []))
    return {
        "to":  list(set(to_recipients)),
        "cc":  list(set(cc_recipients)),
        "bcc": list(set(bcc_recipients)),
    }


# ── EMAIL: Build body and dispatch .docx as attachment ────────────────────────
async def send_report_email(
    filepath,
    start_dt,
    end_dt,
    email_type,
    audience,
):
    recipients = await get_email_users_by_type(
        email_type=email_type, audience=audience
    )

    subject = (
        f"OPC Telemetry Loss of Communication Report — "
        f"{start_dt.strftime('%B %Y')}"
    )

    body = (
        f"Dear Team,\n\n"
        f"Please find attached the OPC Telemetry Loss of Communication Report "
        f"for the period {start_dt.strftime('%B %d, %Y')} – {end_dt.strftime('%B %d, %Y')}.\n\n"
        f"Threshold : 30 minutes\n"
        f"Generated at  : {datetime.now(IST).strftime('%d-%B-%Y %I:%M %p')} IST\n\n"
        f"Regards,\nNovex Automated Reporting"
    )

    ins = await notification_factory.get_notification_module("email")  
    await ins.publish_message(                                         
        subject=subject,                                                
        recipients=recipients["to"],                                    
        cc_recipients=recipients["cc"],                                 
        bcc_recipients=recipients["bcc"],                               
        body=body,                                                      
        attachments=[filepath],   # .docx path from generate_word_report()
        force_send=True,                                                
    )                                                                   
           


# ── ThingsBoard Auth ──────────────────────────────────────────────────────────
async def get_jwt(session):
    url     = f"{THINGSBOARD_URL}/api/auth/login"
    payload = {"username": THINGSBOARD_USERNAME, "password": THINGSBOARD_PASSWORD}
    async with session.post(url, json=payload) as resp:
        if resp.status != 200:
            raise RuntimeError(f"ThingsBoard auth failed: {resp.status}")
        data = await resp.json()
        return data["token"]


def auth_headers(jwt):
    return {"Content-Type": "application/json", "X-Authorization": f"Bearer {jwt}"}


# ── Device helpers ────────────────────────────────────────────────────────────
async def get_all_devices(session, jwt):
    devices, page = [], 0
    while True:
        url = f"{THINGSBOARD_URL}/api/tenant/devices?pageSize={PAGE_SIZE}&page={page}"
        async with session.get(url, headers=auth_headers(jwt)) as resp:
            if resp.status != 200:
                raise RuntimeError(f"Failed to fetch devices: {resp.status}")
            body = await resp.json()
            devices.extend(body.get("data", []))
            if body.get("hasNext", False):
                page += 1
            else:
                break
    return devices


def normalize_location(location):
    return re.sub(r'_[A-Za-z0-9]+$', '', location.strip())


def group_tk_devices_by_location(devices):
    location_map = {}
    for dev in devices:
        name = dev.get("name", "")
        if "@" not in name:
            continue
        local_name, raw_location = name.split("@", 1)
        local_name = local_name.strip()
        location   = normalize_location(raw_location)
        if location.lower() in EXCLUDED_LOCATIONS:
            continue
        is_tk            = local_name.upper().startswith("TK")
        exception_prefix = TANK_DEVICE_EXCEPTIONS.get(location.lower())
        is_exception     = bool(exception_prefix and local_name.upper().startswith(exception_prefix.upper()))
        if (is_tk or is_exception) and location not in location_map:
            location_map[location] = dev
    return location_map


# ── Telemetry helpers ─────────────────────────────────────────────────────────
async def fetch_telemetry(session, jwt, device_id, key, start_ms, end_ms):
    records, window_end = [], end_ms
    while window_end > start_ms:
        url = (
            f"{THINGSBOARD_URL}/api/plugins/telemetry/DEVICE/{device_id}/values/timeseries"
            f"?keys={key}&startTs={start_ms}&endTs={window_end}&limit={PAGE_SIZE}&orderBy=DESC"
        )
        async with session.get(url, headers=auth_headers(jwt)) as resp:
            if resp.status != 200:
                break
            body = await resp.json()
        points = body.get(key, [])
        if not points:
            break
        records.extend(points)
        oldest_ts = points[-1]["ts"]
        if oldest_ts <= start_ms or len(points) < PAGE_SIZE:
            break
        window_end = oldest_ts - 1
    seen   = set()
    unique = [p for p in records if not (p["ts"] in seen or seen.add(p["ts"]))]
    unique.sort(key=lambda x: x["ts"], reverse=True)
    return unique


def find_gaps(records, gap_threshold_min):
    gaps, threshold_ms = [], gap_threshold_min * 60 * 1000
    for i in range(len(records) - 1):
        newer_ts = records[i]["ts"]
        older_ts = records[i + 1]["ts"]
        diff_ms  = newer_ts - older_ts
        if diff_ms > threshold_ms:
            gaps.append({
                "gap_start":    datetime.fromtimestamp(older_ts / 1000, tz=IST),
                "gap_end":      datetime.fromtimestamp(newer_ts / 1000, tz=IST),
                "duration_min": round(diff_ms / 60000, 1),
            })
    return gaps


def get_severity(duration_min):
    hours = duration_min / 60
    if duration_min <= 0:  return "None",     "38761d"
    elif hours < 2:        return "Low",      "bf9000"
    elif hours < 10:       return "Medium",   "b45f06"
    elif hours < 20:       return "High",     "ea9999"
    else:                  return "Critical", "C0392B"


# ── Word Document Generation ──────────────────────────────────────────────────
def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_width(table, width_inches):
    table.width = Inches(width_inches)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(width_inches / len(row.cells))


def generate_word_report(report, location_map, start_dt, end_dt):
    doc = Document()

    style = doc.styles['Normal']
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OPC TELEMETRY LOSS OF COMMUNICATION REPORT")
    run.font.size      = Pt(24)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Report Date: {datetime.now(IST).strftime('%B %d, %Y')} | ")
    info.add_run(f"Period: {start_dt.strftime('%b %d')} – {end_dt.strftime('%b %d, %Y')}\n")
    info.add_run(f"Threshold: 30m | Total Locations: {len(location_map)}")

    doc.add_heading("1. Executive Summary", level=1)

    by_location = defaultdict(list)
    for row in report:
        by_location[row["location"]].append(row)

    total_outages          = len(report)
    locations_with_outages = len(by_location)
    outage_free_locations  = len(location_map) - locations_with_outages
    total_downtime_min     = sum(r["duration_min"] for r in report)
    avg_downtime           = total_downtime_min / len(location_map) if location_map else 0

    summary_text = (
        f"This report covers LOC analysis for {len(location_map)} locations over the period "
        f"{start_dt.strftime('%B %d, %Y')} – {end_dt.strftime('%B %d, %Y')}. "
        f"A Loss of Communication is defined as any interval exceeding 30 minutes without a data point."
    )
    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    set_table_width(table, 7.5)
    headers = ["Monitored", "With LOC's", "LOC-Free", "Total Events", "Total Downtime", "Avg Downtime"]
    vals = [
        str(len(location_map)),
        str(locations_with_outages),
        str(outage_free_locations),
        str(total_outages),
        f"{int(total_downtime_min // 60)}h {int(total_downtime_min % 60)}m",
        f"{int(avg_downtime // 60)}h {int(avg_downtime % 60)}m",
    ]
    for i, (h, v) in enumerate(zip(headers, vals)):
        cell_h = table.cell(0, i)
        cell_h.text = h
        set_cell_background(cell_h, "EBF1F8")
        cell_h.paragraphs[0].runs[0].font.size = Pt(8)
        cell_h.paragraphs[0].runs[0].font.bold = True
        cell_v = table.cell(1, i)
        cell_v.text = v
        cell_v.paragraphs[0].runs[0].font.size = Pt(9)
        cell_v.paragraphs[0].runs[0].font.bold = True

    h2 = doc.add_heading("Severity Classification", level=2)
    h2.paragraph_format.space_before = Pt(12)

    sev_table = doc.add_table(rows=6, cols=3)
    sev_table.style = 'Table Grid'
    set_table_width(sev_table, 7.5)
    for i, h in enumerate(["Level", "Threshold", "Description"]):
        sev_table.cell(0, i).text = h
        set_cell_background(sev_table.cell(0, i), "1F4E78")
        run = sev_table.cell(0, i).paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    sev_data = [
        ("None",     "0 min",      "No LOC detected.",                                           "38761d"),
        ("Low",      "< 2h",       "Minor interruptions. Likely transient connectivity issues.", "bf9000"),
        ("Medium",   "2h – 10h",   "Moderate outages. Requires investigation.",                  "b45f06"),
        ("High",     "10h – 20h",  "Significant data loss. Immediate action recommended.",       "ea9999"),
        ("Critical", "> 20h",      "Severe prolonged outage. Escalation required.",              "ff0000"),
    ]
    for i, (lvl, thr, desc, color) in enumerate(sev_data, 1):
        sev_table.cell(i, 0).text = lvl
        sev_table.cell(i, 1).text = thr
        sev_table.cell(i, 2).text = desc
        set_cell_background(sev_table.cell(i, 0), color)
        if lvl in ["None", "Low", "Medium", "Critical"]:
            sev_table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    h1_2 = doc.add_heading("2. Location-wise Summary", level=1)
    h1_2.paragraph_format.space_before = Pt(18)

    loc_table = doc.add_table(rows=1, cols=4)
    loc_table.style = 'Table Grid'
    set_table_width(loc_table, 7.5)
    for i, h in enumerate(["Location", "LOC's", "Total Downtime", "Severity"]):
        loc_table.cell(0, i).text = h
        set_cell_background(loc_table.cell(0, i), "1F4E78")
        run = loc_table.cell(0, i).paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    for location in sorted(location_map.keys()):
        rows      = by_location[location]
        total_min = sum(r["duration_min"] for r in rows)
        sev_lvl, sev_color = get_severity(total_min)
        row_cells = loc_table.add_row().cells
        row_cells[0].text = location
        row_cells[1].text = str(len(rows)) if rows else "—"
        row_cells[2].text = f"{int(total_min // 60)}h {int(total_min % 60)}m" if total_min > 0 else "—"
        row_cells[3].text = sev_lvl
        set_cell_background(row_cells[3], sev_color)
        if sev_lvl in ["None", "Low", "Medium", "Critical"]:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    h1_3 = doc.add_heading("3. Detailed LOC Logs by Location", level=1)
    h1_3.paragraph_format.space_before = Pt(18)
    doc.add_paragraph("All timestamps are in Indian Standard Time (IST, UTC+5:30).")

    for location, rows in sorted(by_location.items()):
        total_min = sum(r["duration_min"] for r in rows)
        h2_loc = doc.add_heading(
            f"{location} — {len(rows)} LOC, {int(total_min // 60)}h {int(total_min % 60)}m total",
            level=2,
        )
        h2_loc.paragraph_format.space_before = Pt(12)
        sev_lvl, _ = get_severity(total_min)
        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        p.add_run(sev_lvl)

        log_table = doc.add_table(rows=1, cols=3)
        log_table.style = 'Table Grid'
        set_table_width(log_table, 7.5)
        for i, h in enumerate(["Last Record (IST)", "Data Resumed (IST)", "Duration"]):
            log_table.cell(0, i).text = h
            set_cell_background(log_table.cell(0, i), "1F4E78")
            run = log_table.cell(0, i).paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
        for r in sorted(rows, key=lambda x: x["gap_start"]):
            row_cells = log_table.add_row().cells
            row_cells[0].text = r["gap_start"].strftime("%Y-%m-%d %H:%M:%S")
            row_cells[1].text = r["gap_end"].strftime("%Y-%m-%d %H:%M:%S")
            row_cells[2].text = f"{int(r['duration_min'] // 60)}h {int(r['duration_min'] % 60)}m"

    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
    os.makedirs(output_dir, exist_ok=True)
    filename = os.path.join(output_dir, f"OPC_Connectivity_Outage_Report_{start_dt.strftime('%B%Y')}.docx")
    doc.save(filename)
    return filename


# ── Main ──────────────────────────────────────────────────────────────────────
async def main(email_type: str = "loc_report", audience: str = "TAS"):
    now_ist             = datetime.now(IST)
    current_month_start = now_ist.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    end_dt              = current_month_start - timedelta(microseconds=1)
    start_dt            = end_dt.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    start_ms            = int(start_dt.timestamp() * 1000)
    end_ms              = int(end_dt.timestamp() * 1000)

    async with aiohttp.ClientSession() as session:
        
        jwt = await get_jwt(session)

        
        devices      = await get_all_devices(session, jwt)
        location_map = group_tk_devices_by_location(devices)
        if not location_map:
            print("  No TK-prefixed devices found. Exiting.")
            return

        
        report = []
        for location, device in sorted(location_map.items()):
            device_id = device["id"]["id"]
            records   = await fetch_telemetry(session, jwt, device_id, TELEMETRY_KEY, start_ms, end_ms)
            # print(f"  → {location:20s} | {len(records):5d} pts")
            if not records:
                report.append({
                    "location":     location,
                    "gap_start":    start_dt,
                    "gap_end":      end_dt,
                    "duration_min": round((end_ms - start_ms) / 60000, 1),
                })
                continue
            for g in find_gaps(records, GAP_THRESHOLD_MIN):
                report.append({"location": location, **g})

        
        word_file = generate_word_report(report, location_map, start_dt, end_dt)
        print(f"      Saved: {word_file}")

        # ── EMAIL: send .docx as attachment ───────────────────────────────────
                                   
        await send_report_email(                                     
            filepath=word_file,                                      
            start_dt=start_dt,                                       
            end_dt=end_dt,                                           
            email_type=email_type,                                   
            audience=audience,                                       
        )                                                            
                                               


if __name__ == "__main__":
    _email_type = sys.argv[1] if len(sys.argv) > 1 else "tas_loc_report"
    _audience   = sys.argv[2] if len(sys.argv) > 2 else "tas_loc_report_employee"
    asyncio.run(main(email_type=_email_type, audience=_audience))